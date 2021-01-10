from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
import os
import datetime
import threading
import speech_recognition as sr
import sys

# Do I need this?
sys.path.insert(1, "../..")
import speech_test
import months_gui_test

def start_setup():
    x = datetime.datetime.now()
    month = x.strftime("%B")
    year = x.strftime("%Y")
    date = x.strftime("%B %#d, %Y")

    path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\%s_%s.docx" % (month,year)
 
    # speechChecker = threading.Thread(target=start_speech,args=(path,))
    # speechChecker.start()   
    
    if os.path.isfile(path) == True:
        document = Document(path)
        all_paras = document.paragraphs     
        # search through the dates of the document and compare it to today's date
        for para in all_paras:
            # If find a date matching today, then pull up user interface for selecting which month to open
            if date in para.text:
                root = months_gui_test.root
                app = months_gui_test.Application(master=root)
                app.mainloop()
                return
        # If do not find a date matching today, then add today's date for new entry                
        add_date(document, date)
    else:
        document = Document()
        create_doc(document, date, month)
        
    document.save(path)
    os.startfile(path, 'open')

def add_date(document, date):
        paragraph = document.add_paragraph('')

        run = paragraph.add_run(date)
        run.bold = True
        run.italic = True

        font = run.font
        font.name = 'Cambria'
        font.size = Pt(11)

        paragraph_2 = document.add_paragraph('')
        
        run2 = paragraph_2.add_run('Entry: ')

        font2 = run2.font
        font2.name = 'Cambria'
        font2.size = Pt(11)

def create_doc(document, date, month): 
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = title.add_run(month)

        font = title_text.font
        font.name = 'Cambria'
        font.size = Pt(28)

        paragraph = document.add_paragraph('')

        run = paragraph.add_run(date)
        run.bold = True
        run.italic = True

        font = run.font
        font.name = 'Cambria'
        font.size = Pt(11)

        paragraph_2 = document.add_paragraph('')
        
        run2 = paragraph_2.add_run('Entry: ')

        font2 = run2.font
        font2.name = 'Cambria'
        font2.size = Pt(11)

def start_speech(path):
    closingWords = ['thank you', 'thanks', 'later', 'bye bye', 'bye', 'good night']
    # create recognizer and mic instances
    recognizer = sr.Recognizer()
    microphone = sr.Microphone()
    listening = True
    
    while listening:
    
        while True:
            words = speech_test.recognize_speech_from_mic(recognizer, microphone)
            if words:
                break
        
        for word in closingWords:
            if word in words:
                listening = False
                os.system("TASKKILL /IM winword.exe")


start_setup()

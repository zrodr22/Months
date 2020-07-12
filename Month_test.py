from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
import os
import datetime
import threading

# Add a thread here that will run the speech recognition. Call a function in thread to close and save doc

# Split this up in to different functions?
def start_setup():
    x = datetime.datetime.now()
    month = x.strftime("%B")
    year = x.strftime("%Y")
    date = x.strftime("%B %#d, %Y")

    path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\%s_%s.docx" % (month,year)
    
    #path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\Feb_Q.docx"
    if os.path.isfile(path) == True:
        # search through the dates of the document and compare it to today's date

        # If do not find a date matching today, then add today's date for new entry
        
        #path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\%January_Q.docx"
        print(path)
        document = Document(path)
    
        paragraph = document.add_paragraph('')

        run = paragraph.add_run(date)
        run.bold = True
        run.italic = True

        font = run.font
        font.name = 'Cambria'
        font.size = Pt(11)

        paragraph_2 = document.add_paragraph('')
        
        run2 = paragraph_2.add_run('Entry: ')

        font2 = run2.font
        font2.name = 'Cambria'
        font2.size = Pt(11)

        document.save(path)

        os.startfile(path, 'open')

        # If find a date matching today, then pull up user interface for selecting which month to open

            # Figure out how to select from existing months to open
    else:
        document = Document()
        path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\%s_%s.docx" % (month,year)
        #path = "C:\\Users\\zrodr\\Documents\\ThoughtJourneys\\Months\\%s_%s.docx" % ("zach",year)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = title.add_run(month)

        font = title_text.font
        font.name = 'Cambria'
        font.size = Pt(28)

        paragraph = document.add_paragraph('')

        run = paragraph.add_run(date)
        run.bold = True
        run.italic = True

        font = run.font
        font.name = 'Cambria'
        font.size = Pt(11)

        paragraph_2 = document.add_paragraph('')
        
        run2 = paragraph_2.add_run('Entry: ')

        font2 = run2.font
        font2.name = 'Cambria'
        font2.size = Pt(11)

        document.save(path)

        os.startfile(path, 'open')

start_setup()
